"""
BCD Portal Anonymizer — unified de-identification pipeline for the
breast-cancer-image-dataset bucket.

Combines:
  - pydicom metadata scrubbing (DICOM/BIN/DICOM files)
  - privacy_filter PDFRedactor (native-text + scanned PDF redaction)
  - python-docx + SafeHarborDetector (DOCX text redaction)
  - EkaCare document-pii-redactor (image PII: consent, biopsy, ultrasound)

Design principles:
  - DICOM: metadata-only, pixel data UNTOUCHED, write_like_original=True
  - Documents: strip PII text/images, preserve clinical findings
  - Images: detect + redact PII, strip EXIF, preserve format/quality
  - Clinical allowlist: BIRADS, breast density, lesion findings are NEVER redacted
"""

from __future__ import annotations

import io
import logging
import os
import re
import shutil
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any

import pydicom
from pydicom.dataset import Dataset
from pydicom.uid import generate_uid

logger = logging.getLogger("bcd_anonymizer")


class FileCategory(str, Enum):
    DICOM = "dicom"
    PDF = "pdf"
    DOCX = "docx"
    IMAGE = "image"
    PASSTHROUGH = "passthrough"


@dataclass
class AnonymizationResult:
    input_path: str
    output_path: str
    category: FileCategory
    success: bool
    phi_found: list[dict] = field(default_factory=list)
    phi_removed: int = 0
    error: str | None = None
    file_size_before: int = 0
    file_size_after: int = 0
    copied_unanonymized: bool = False

_MIN_DICOM_SIZE = 1024


# ── DICOM PII tags to scrub ─────────────────────────────────────────────────
# Only fields that contain patient/staff/institution identifiers.
# Imaging parameters, modality info, geometry, etc. are preserved.
DICOM_PII_TAGS = {
    # Patient identifiers
    "PatientName",
    "PatientID",
    "IssuerOfPatientID",
    "OtherPatientIDs",
    "OtherPatientNames",
    "PatientBirthDate",
    "PatientBirthTime",
    "PatientAge",
    "PatientAddress",
    "PatientTelephoneNumbers",
    "EthnicGroup",
    "PatientComments",
    "AdditionalPatientHistory",
    # Institutional identifiers
    "AccessionNumber",
    "InstitutionName",
    "InstitutionAddress",
    "StationName",
    "DeviceSerialNumber",
    "InstitutionalDepartmentName",
    # Personnel
    "ReferringPhysicianName",
    "RequestingPhysician",
    "OperatorsName",
    "PerformingPhysicianName",
    "NameOfPhysiciansReadingStudy",
    "PhysiciansOfRecord",
    "ResponsiblePerson",
    # Study-level identifiers (can embed PII in free-text)
    "StudyID",
    "RequestedProcedureDescription",
}

# VRs that are not free-text — cannot hold "REDACTED" as a replacement.
_NON_TEXT_VRS = {"DA", "DT", "TM", "AS", "IS", "DS", "FL", "FD",
                 "SL", "SS", "UL", "US", "UI", "AT", "OB", "OW",
                 "OF", "OD", "UN", "SQ"}

# VRs that can hold text (potential PII in private tags).
_TEXT_VRS = {"LO", "PN", "SH", "LT", "ST", "UT", "CS", "AE"}

# ── Breast cancer clinical allowlist ─────────────────────────────────────────
# These terms must NEVER be redacted from documents or images.
CLINICAL_ALLOWLIST = {
    # BIRADS
    "bi-rads", "birads", "bi rads",
    "birads 0", "birads 1", "birads 2", "birads 3",
    "birads 4", "birads 4a", "birads 4b", "birads 4c",
    "birads 5", "birads 6",
    "bi-rads 0", "bi-rads 1", "bi-rads 2", "bi-rads 3",
    "bi-rads 4", "bi-rads 4a", "bi-rads 4b", "bi-rads 4c",
    "bi-rads 5", "bi-rads 6",
    "category 0", "category 1", "category 2", "category 3",
    "category 4", "category 5", "category 6",

    # Breast density
    "breast density", "heterogeneously dense", "scattered fibroglandular",
    "extremely dense", "fatty", "dense breast", "acr a", "acr b", "acr c", "acr d",
    "density a", "density b", "density c", "density d",

    # Mammographic findings
    "mass", "calcification", "calcifications", "microcalcification",
    "microcalcifications", "architectural distortion", "asymmetric density",
    "focal asymmetry", "global asymmetry", "developing asymmetry",
    "skin thickening", "skin retraction", "nipple retraction",
    "axillary lymphadenopathy", "intramammary lymph node",

    # Lesion descriptors
    "spiculated", "lobulated", "irregular", "oval", "round",
    "circumscribed", "obscured", "indistinct", "microlobulated",
    "pleomorphic", "amorphous", "coarse", "punctate",
    "linear", "segmental", "regional", "diffuse", "grouped", "clustered",

    # Anatomical
    "right breast", "left breast", "bilateral", "unilateral",
    "upper outer quadrant", "upper inner quadrant",
    "lower outer quadrant", "lower inner quadrant",
    "retroareolar", "subareolar", "axillary", "axilla",
    "mediolateral oblique", "mlo", "craniocaudal", "cc",
    "mammogram", "mammography", "mammographic",
    "ultrasound", "sonography", "sonographic",
    "biopsy", "fnac", "core biopsy", "excision biopsy",

    # General clinical
    "benign", "malignant", "suspicious", "probably benign",
    "highly suggestive of malignancy", "incomplete",
    "negative", "normal", "abnormal",
    "invasive ductal carcinoma", "invasive lobular carcinoma",
    "ductal carcinoma in situ", "dcis", "idc", "ilc",
    "fibroadenoma", "cyst", "phyllodes",
    "carcinoma", "adenocarcinoma", "lymphoma",
    "neoplasm", "tumor", "tumour", "lesion",
    "metastasis", "metastatic",
    "estrogen receptor", "progesterone receptor", "her2",
    "ki-67", "ki67",
    "sentinel lymph node", "mastectomy", "lumpectomy",
    "chemotherapy", "radiation therapy", "hormonal therapy",

    # Report structure
    "impression", "findings", "conclusion", "recommendation",
    "clinical history", "indication", "technique",
    "comparison", "prior study",

    # Mammographic view labels (burned-in annotations on images)
    "l cc", "r cc", "l mlo", "r mlo",
    "lcc", "rcc", "lmlo", "rmlo",
    "l lat", "r lat", "lat",
    "l xccl", "r xccl",
}


def _is_clinical_term(text: str) -> bool:
    t = text.strip().lower()
    return t in CLINICAL_ALLOWLIST


# ═══════════════════════════════════════════════════════════════════════════════
#  FILE CATEGORY DETECTION
# ═══════════════════════════════════════════════════════════════════════════════

def detect_category(path: Path) -> FileCategory:
    suffix = path.suffix.lower()

    if suffix in {".dcm", ".dicom"}:
        return FileCategory.DICOM

    if suffix == ".bin":
        if _is_dicom_file(path):
            return FileCategory.DICOM
        return FileCategory.PASSTHROUGH

    if suffix == ".pdf":
        return FileCategory.PDF

    if suffix == ".docx":
        return FileCategory.DOCX

    if suffix in {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp",
                  ".gif", ".webp", ".ico", ".svg"}:
        return FileCategory.IMAGE

    if suffix in {".zip", ".xlsx", ".xls", ".csv", ".txt", ".json", ".xml",
                  ".avi", ".mp4", ".mov", ".doc"}:
        return FileCategory.PASSTHROUGH

    # Try DICOM magic bytes for unknown extensions
    if _is_dicom_file(path):
        return FileCategory.DICOM

    return FileCategory.PASSTHROUGH


def _is_dicom_file(path: Path) -> bool:
    try:
        if not path.exists() or path.stat().st_size < _MIN_DICOM_SIZE:
            return False
        with open(path, "rb") as f:
            f.seek(128)
            magic = f.read(4)
            if magic == b"DICM":
                return True
            f.seek(0)
            header = f.read(256)
            if b"\x08\x00\x05\x00" in header or b"\x08\x00\x10\x00" in header:
                return True
    except Exception:
        return False
    try:
        pydicom.dcmread(str(path), force=True, stop_before_pixels=True)
        return True
    except Exception:
        return False


# ═══════════════════════════════════════════════════════════════════════════════
#  DICOM ANONYMIZATION — metadata only, pixel data untouched
# ═══════════════════════════════════════════════════════════════════════════════

def anonymize_dicom(
    input_path: Path,
    output_path: Path,
    subject_id: str = "ANONYMOUS",
) -> AnonymizationResult:
    file_size = input_path.stat().st_size if input_path.exists() else 0
    result = AnonymizationResult(
        input_path=str(input_path),
        output_path=str(output_path),
        category=FileCategory.DICOM,
        success=False,
        file_size_before=file_size,
    )

    if file_size < _MIN_DICOM_SIZE:
        result.error = f"Skipped: file too small ({file_size} bytes), likely corrupt or truncated"
        return result

    try:
        ds = pydicom.dcmread(str(input_path), force=True)
    except Exception as e:
        result.error = f"Failed to read DICOM: {e}"
        return result

    if not hasattr(ds, "PixelData") and not hasattr(ds, "Rows"):
        result.error = f"Skipped: DICOM has no pixel data or image dimensions, likely truncated"
        return result

    original_ts = None
    if hasattr(ds, "file_meta") and hasattr(ds.file_meta, "TransferSyntaxUID"):
        original_ts = ds.file_meta.TransferSyntaxUID

    if original_ts is None:
        from pydicom.uid import ImplicitVRLittleEndian
        original_ts = ImplicitVRLittleEndian
        if not hasattr(ds, "file_meta") or ds.file_meta is None:
            ds.file_meta = Dataset()
        ds.file_meta.TransferSyntaxUID = original_ts

    # ── Scrub PII tags ────────────────────────────────────────────────────
    for tag_name in DICOM_PII_TAGS:
        if not hasattr(ds, tag_name):
            continue
        try:
            elem = ds.data_element(tag_name)
            if elem is None:
                continue
            original_val = str(elem.value).strip()
            if not original_val:
                continue

            result.phi_found.append({
                "tag": tag_name,
                "original_value": original_val,
                "source": "dicom_metadata",
            })

            vr = elem.VR if elem else ""
            if vr in _NON_TEXT_VRS:
                elem.value = ""
            elif tag_name == "PatientID":
                elem.value = subject_id
            elif tag_name == "PatientName":
                elem.value = "ANONYMOUS"
            elif vr == "PN":
                elem.value = "ANONYMOUS"
            else:
                elem.value = ""

            result.phi_removed += 1
        except Exception as e:
            logger.warning("Failed to clean DICOM tag %s: %s", tag_name, e)

    # ── Remove ALL private tags (DICOM PS3.15 de-identification) ────────
    # Private tags are vendor-specific extensions (e.g. FDMS, Carestream).
    # No standard imaging parameter lives in a private tag — all critical
    # data (Rows, Columns, BitsAllocated, PixelSpacing, etc.) is in public
    # groups.
    private_removed = 0
    for elem in list(ds):
        if elem.tag.is_private:
            vr = getattr(elem, "VR", "")
            val = str(elem.value).strip()
            if vr in _TEXT_VRS and val and val != "None":
                result.phi_found.append({
                    "tag": f"Private({elem.tag})",
                    "original_value": val[:100],
                    "source": "dicom_private_tag",
                })
            private_removed += 1

    ds.remove_private_tags()
    result.phi_removed += private_removed

    # ── Mark as de-identified ─────────────────────────────────────────────
    ds.PatientIdentityRemoved = "YES"
    ds.DeidentificationMethod = "BCD Portal Anonymizer | metadata scrub"

    # ── Ensure valid file meta ────────────────────────────────────────────
    if not hasattr(ds, "file_meta") or ds.file_meta is None:
        ds.file_meta = Dataset()

    ds.file_meta.TransferSyntaxUID = original_ts

    if not getattr(ds.file_meta, "MediaStorageSOPClassUID", None):
        ds.file_meta.MediaStorageSOPClassUID = getattr(
            ds, "SOPClassUID", "1.2.840.10008.5.1.4.1.1.7"
        )
    if not getattr(ds.file_meta, "MediaStorageSOPInstanceUID", None):
        ds.file_meta.MediaStorageSOPInstanceUID = getattr(
            ds, "SOPInstanceUID", generate_uid()
        )

    # ── Save with original transfer syntax ────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        ds.save_as(str(output_path), write_like_original=True)
    except Exception:
        ds.save_as(str(output_path), write_like_original=False)
        logger.warning(
            "write_like_original failed for %s, fell back to write_like_original=False",
            input_path.name,
        )

    result.file_size_after = output_path.stat().st_size
    result.success = True
    return result


# ═══════════════════════════════════════════════════════════════════════════════
#  PDF ANONYMIZATION — native text Safe Harbor + scanned OCR + QR removal
# ═══════════════════════════════════════════════════════════════════════════════

def anonymize_pdf(
    input_path: Path,
    output_path: Path,
) -> AnonymizationResult:
    result = AnonymizationResult(
        input_path=str(input_path),
        output_path=str(output_path),
        category=FileCategory.PDF,
        success=False,
        file_size_before=input_path.stat().st_size,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        from privacy_filter.redactors.pdf_redactor import PDFRedactor
        redactor = PDFRedactor()
        entities, counts = redactor.redact(input_path, output_path)

        for e in entities:
            result.phi_found.append({
                "category": e.get("entity_group", "UNKNOWN"),
                "text": e.get("word", ""),
                "page": e.get("bbox", {}).get("page", 0),
                "source": "pdf_safe_harbor",
            })

        result.phi_removed = len(entities)
        result.file_size_after = output_path.stat().st_size
        result.success = True

    except Exception as e:
        result.error = f"PDF anonymization failed: {e}"
        logger.error("PDF anonymization failed for %s: %s", input_path.name, e)
        try:
            shutil.copy2(str(input_path), str(output_path))
            result.file_size_after = output_path.stat().st_size
            result.copied_unanonymized = True
            result.error += " (copied as-is — may still contain PII)"
        except Exception:
            pass

    return result


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCX ANONYMIZATION — text-level Safe Harbor redaction in-place
# ═══════════════════════════════════════════════════════════════════════════════

def anonymize_docx(
    input_path: Path,
    output_path: Path,
) -> AnonymizationResult:
    result = AnonymizationResult(
        input_path=str(input_path),
        output_path=str(output_path),
        category=FileCategory.DOCX,
        success=False,
        file_size_before=input_path.stat().st_size,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document
        from privacy_filter.detectors.safe_harbor_detector import SafeHarborDetector

        doc = Document(str(input_path))
        detector = SafeHarborDetector()

        total_redacted = 0

        # Process paragraphs
        for para in doc.paragraphs:
            total_redacted += _redact_docx_paragraph(para, detector, result)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        total_redacted += _redact_docx_paragraph(
                            para, detector, result
                        )

        # Process headers and footers
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                if header_footer is not None:
                    for para in header_footer.paragraphs:
                        total_redacted += _redact_docx_paragraph(
                            para, detector, result
                        )

        result.phi_removed = total_redacted

        # Remove document metadata
        core = doc.core_properties
        _docx_meta_fields = [
            "author", "last_modified_by", "comments", "category",
            "content_status", "identifier", "keywords", "subject",
            "title", "description",
        ]
        for field_name in _docx_meta_fields:
            try:
                current = getattr(core, field_name, None)
                if current:
                    result.phi_found.append({
                        "category": f"DOCX_META:{field_name}",
                        "text": str(current)[:100],
                        "source": "docx_metadata",
                    })
                    setattr(core, field_name, "")
                    result.phi_removed += 1
            except Exception:
                pass

        doc.save(str(output_path))
        result.file_size_after = output_path.stat().st_size
        result.success = True

    except Exception as e:
        result.error = f"DOCX anonymization failed: {e}"
        logger.error("DOCX anonymization failed for %s: %s", input_path.name, e)
        try:
            shutil.copy2(str(input_path), str(output_path))
            result.file_size_after = output_path.stat().st_size
            result.copied_unanonymized = True
            result.error += " (copied as-is — may still contain PII)"
        except Exception:
            pass

    return result


def _redact_docx_paragraph(para, detector, result: AnonymizationResult) -> int:
    text = para.text
    if not text or not text.strip():
        return 0

    spans = detector.detect_spans(text)
    if not spans:
        return 0

    # Check clinical allowlist
    spans = [s for s in spans if not _is_clinical_term(s.text)]
    if not spans:
        return 0

    for span in spans:
        result.phi_found.append({
            "category": span.label,
            "text": span.text[:80],
            "source": "docx_safe_harbor",
        })

    # Build redacted text by replacing PII spans with [REDACTED]
    redacted = _apply_span_redactions(text, spans)

    # Replace paragraph text while preserving formatting.
    # We clear all runs and set the first run's text.
    if para.runs:
        for i, run in enumerate(para.runs):
            if i == 0:
                run.text = redacted
            else:
                run.text = ""
    else:
        para.text = redacted

    return len(spans)


def _apply_span_redactions(text: str, spans) -> str:
    # Sort spans by start position, descending (so we can replace right-to-left)
    sorted_spans = sorted(spans, key=lambda s: s.start, reverse=True)
    chars = list(text)
    for span in sorted_spans:
        chars[span.start:span.end] = list("[REDACTED]")
    return "".join(chars)


# ═══════════════════════════════════════════════════════════════════════════════
#  IMAGE ANONYMIZATION — EkaCare PII detection + EXIF strip
# ═══════════════════════════════════════════════════════════════════════════════

def anonymize_image(
    input_path: Path,
    output_path: Path,
) -> AnonymizationResult:
    result = AnonymizationResult(
        input_path=str(input_path),
        output_path=str(output_path),
        category=FileCategory.IMAGE,
        success=False,
        file_size_before=input_path.stat().st_size,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        from PIL import Image

        img = Image.open(input_path)
        original_format = img.format or input_path.suffix.lstrip(".").upper()
        original_mode = img.mode

        pii_entities = []

        # ── Try EkaCare PII detection ─────────────────────────────────────
        try:
            from document_pii_redactor import ImagePIIRedactor
            redactor = _get_ekacare_redactor()

            try:
                entities = redactor.detect(str(input_path))
            except Exception as e:
                logger.warning("EkaCare detect() failed for %s: %s", input_path.name, e)
                entities = []

            # Filter out clinical terms from entities
            filtered = []
            for e in entities:
                e_text = getattr(e, "text", None) or ""
                e_cat = getattr(e, "category", None) or ""
                if _is_clinical_term(e_text):
                    continue
                if e_cat in {
                    "test_result_value", "test_result_name",
                    "diagnosis", "medical_condition",
                    "logo",
                }:
                    continue
                filtered.append(e)

            if filtered:
                for e in filtered:
                    result.phi_found.append({
                        "category": getattr(e, "category", None) or "unknown",
                        "text": (getattr(e, "text", None) or "")[:80],
                        "score": getattr(e, "score", 0),
                        "source": "ekacare",
                    })

                try:
                    redacted = redactor.redact(str(input_path), filtered, mode="solid")
                    img = redacted
                    pii_entities = filtered
                except Exception as e:
                    logger.warning("EkaCare redact() failed for %s: %s", input_path.name, e)
        except ImportError:
            logger.info("EkaCare not available, using SafeHarbor fallback for %s", input_path.name)

        # ── Fallback: SafeHarbor OCR scan for document-like images ────────
        if not pii_entities:
            try:
                safe_harbor_entities = _safe_harbor_image_scan(input_path)
                if safe_harbor_entities:
                    import cv2
                    import numpy as np
                    img_arr = np.array(img.convert("RGB"))
                    img_arr = cv2.cvtColor(img_arr, cv2.COLOR_RGB2BGR)
                    for e in safe_harbor_entities:
                        bbox = e.get("bbox")
                        if bbox:
                            x1, y1 = int(bbox["x1"]), int(bbox["y1"])
                            x2, y2 = int(bbox["x2"]), int(bbox["y2"])
                            cv2.rectangle(img_arr, (x1, y1), (x2, y2), (0, 0, 0), -1)
                        result.phi_found.append({
                            "category": e.get("label", "UNKNOWN"),
                            "text": e.get("text", "")[:80],
                            "source": "safe_harbor_ocr",
                        })
                    img_arr = cv2.cvtColor(img_arr, cv2.COLOR_BGR2RGB)
                    img = Image.fromarray(img_arr)
            except Exception as e:
                logger.warning("SafeHarbor image scan failed for %s: %s", input_path.name, e)

        result.phi_removed = len(result.phi_found)

        # ── Check if original has EXIF that needs stripping ──────────────
        orig_img = Image.open(input_path)
        has_exif = bool(orig_img.info.get("exif"))

        # If no PII found and no EXIF, copy as-is (preserves exact quality/size)
        if not result.phi_found and not has_exif:
            shutil.copy2(str(input_path), str(output_path))
            result.file_size_after = output_path.stat().st_size
            result.success = True
        else:
            # Re-save with redactions and/or without EXIF
            if hasattr(img, "info"):
                img.info.pop("exif", None)
                img.info.pop("icc_profile", None)

            save_kwargs: dict[str, Any] = {}
            fmt = original_format
            if fmt in ("JPEG", "JPG"):
                fmt = "JPEG"
                if img.mode not in ("RGB", "L"):
                    img = img.convert("RGB")
                save_kwargs["quality"] = 95
            elif fmt == "PNG":
                pass
            elif fmt == "TIFF":
                pass
            elif fmt == "BMP":
                if img.mode == "RGBA":
                    img = img.convert("RGB")
            elif fmt == "WEBP":
                save_kwargs["quality"] = 95
            elif fmt == "GIF":
                pass
            else:
                fmt = "PNG"

            img.save(str(output_path), format=fmt, **save_kwargs)
            result.file_size_after = output_path.stat().st_size
            result.success = True

    except Exception as e:
        result.error = f"Image anonymization failed: {e}"
        logger.error("Image anonymization failed for %s: %s", input_path.name, e)
        try:
            shutil.copy2(str(input_path), str(output_path))
            result.file_size_after = output_path.stat().st_size
            result.copied_unanonymized = True
            result.error += " (copied as-is — may still contain PII)"
        except Exception:
            pass

    return result


# EkaCare redactor singleton (model loading is expensive ~5s)
_ekacare_instance = None


def _get_ekacare_redactor():
    global _ekacare_instance
    if _ekacare_instance is None:
        from document_pii_redactor import ImagePIIRedactor
        _ekacare_instance = ImagePIIRedactor("ekacare/document-pii-redactor")
    return _ekacare_instance


def _safe_harbor_image_scan(image_path: Path) -> list[dict]:
    import cv2
    import numpy as np
    import pytesseract
    from privacy_filter.detectors.safe_harbor_detector import SafeHarborDetector

    img = cv2.imread(str(image_path))
    if img is None:
        return []

    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    data = pytesseract.image_to_data(
        gray, config="--oem 3 --psm 6", output_type=pytesseract.Output.DICT,
    )

    lines: dict[tuple, list] = {}
    n = len(data["text"])
    for i in range(n):
        text = str(data["text"][i]).strip()
        if not text:
            continue
        try:
            conf = float(data["conf"][i])
        except (ValueError, TypeError):
            conf = 0
        if conf < 10:
            continue
        key = (int(data["block_num"][i]), int(data["line_num"][i]))
        lines.setdefault(key, []).append({
            "text": text,
            "x": int(data["left"][i]),
            "y": int(data["top"][i]),
            "w": int(data["width"][i]),
            "h": int(data["height"][i]),
        })

    detector = SafeHarborDetector()
    entities: list[dict] = []

    for key, words in lines.items():
        words.sort(key=lambda w: w["x"])
        line_text = ""
        offsets = []
        for w in words:
            start = len(line_text)
            line_text += w["text"]
            offsets.append((start, len(line_text), w))
            line_text += " "

        for span in detector.detect_spans(line_text):
            if _is_clinical_term(span.text):
                continue
            x1, y1, x2, y2 = None, None, None, None
            for cs, ce, w in offsets:
                if ce <= span.start or cs >= span.end:
                    continue
                wx1, wy1 = w["x"], w["y"]
                wx2, wy2 = w["x"] + w["w"], w["y"] + w["h"]
                if x1 is None:
                    x1, y1, x2, y2 = wx1, wy1, wx2, wy2
                else:
                    x1, y1 = min(x1, wx1), min(y1, wy1)
                    x2, y2 = max(x2, wx2), max(y2, wy2)
            if x1 is not None:
                pad = 4
                entities.append({
                    "label": span.label,
                    "text": span.text,
                    "bbox": {
                        "x1": max(0, x1 - pad),
                        "y1": max(0, y1 - pad),
                        "x2": x2 + pad,
                        "y2": y2 + pad,
                    },
                })

    return entities


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN ENTRY POINT — route files to the correct anonymizer
# ═══════════════════════════════════════════════════════════════════════════════

def anonymize_file(
    input_path: str | Path,
    output_path: str | Path,
    subject_id: str = "ANONYMOUS",
) -> AnonymizationResult:
    input_path = Path(input_path)
    output_path = Path(output_path)

    if not input_path.exists():
        return AnonymizationResult(
            input_path=str(input_path),
            output_path=str(output_path),
            category=FileCategory.PASSTHROUGH,
            success=False,
            error=f"File not found: {input_path.name}",
        )

    file_size = input_path.stat().st_size
    if file_size == 0:
        return AnonymizationResult(
            input_path=str(input_path),
            output_path=str(output_path),
            category=FileCategory.PASSTHROUGH,
            success=False,
            file_size_before=0,
            error=f"Skipped: zero-byte file ({input_path.name})",
        )

    category = detect_category(input_path)
    logger.info("Anonymizing %s [%s]", input_path.name, category.value)

    if category == FileCategory.DICOM:
        return anonymize_dicom(input_path, output_path, subject_id)

    if category == FileCategory.PDF:
        return anonymize_pdf(input_path, output_path)

    if category == FileCategory.DOCX:
        return anonymize_docx(input_path, output_path)

    if category == FileCategory.IMAGE:
        return anonymize_image(input_path, output_path)

    # Passthrough — copy as-is (zip, xlsx, unknown formats)
    result = AnonymizationResult(
        input_path=str(input_path),
        output_path=str(output_path),
        category=FileCategory.PASSTHROUGH,
        success=True,
        file_size_before=file_size,
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(str(input_path), str(output_path))
    result.file_size_after = output_path.stat().st_size
    return result


def anonymize_batch(
    file_pairs: list[tuple[str | Path, str | Path]],
    subject_id: str = "ANONYMOUS",
) -> list[AnonymizationResult]:
    results = []
    for input_path, output_path in file_pairs:
        try:
            r = anonymize_file(input_path, output_path, subject_id)
            results.append(r)
        except Exception as e:
            results.append(AnonymizationResult(
                input_path=str(input_path),
                output_path=str(output_path),
                category=FileCategory.PASSTHROUGH,
                success=False,
                error=str(e),
            ))
    return results
